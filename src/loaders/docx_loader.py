"""
DOCX Document Loader.

Uses python-docx for text extraction from Microsoft Word documents.
"""

from pathlib import Path

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from src.loaders.base import DocumentLoader, RawDocument


class DOCXLoader(DocumentLoader):
    """
    Loader for Microsoft Word (.docx) documents.
    
    Extracts text content and preserves heading structure.
    
    Example:
        loader = DOCXLoader()
        doc = loader.load(Path("document.docx"))
    """
    
    supported_extensions = ["docx"]
    
    def load(self, path: Path) -> RawDocument:
        """Load and extract text from DOCX."""
        if DocxDocument is None:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )
        
        path = self._validate_path(path)
        
        # Open document
        doc = DocxDocument(str(path))
        
        # Extract paragraphs and track headings
        paragraphs = []
        raw_sections = []
        position = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if text:
                paragraphs.append(text)
                
                # Check if this is a heading
                if para.style and para.style.name.startswith('Heading'):
                    try:
                        level = int(para.style.name.replace('Heading ', ''))
                    except ValueError:
                        level = 1
                    
                    raw_sections.append({
                        "title": text,
                        "position": position,
                        "level": level,
                    })
                
                position += len(text) + 2  # +2 for paragraph break
        
        content = "\n\n".join(paragraphs)
        
        # Get file stats
        file_stats = path.stat()
        
        return RawDocument(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type="docx",
            file_size=file_stats.st_size,
            content=content,
            page_count=None,  # DOCX doesn't have fixed pages
            word_count=self._count_words(content),
            raw_sections=raw_sections,
        )
